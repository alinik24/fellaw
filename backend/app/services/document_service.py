"""
Document processing service.

Handles text extraction from PDF/DOCX files and AI-powered document analysis.
"""
from __future__ import annotations

import json
import os
import re
import unicodedata
import uuid
from pathlib import Path
from typing import Any

import aiofiles
import structlog
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.evidence import Document
from app.services.ai_service import chat_completion

log = structlog.get_logger(__name__)

# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------


async def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract plain text from a PDF file using pypdf.

    Iterates over all pages and concatenates extracted text.
    Returns an empty string if extraction fails.
    """
    try:
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        parts: list[str] = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
        full_text = "\n\n".join(parts)
        log.info(
            "extract_text_from_pdf.done",
            file=file_path,
            pages=len(reader.pages),
            chars=len(full_text),
        )
        return full_text
    except Exception as exc:
        log.error("extract_text_from_pdf.failed", file=file_path, error=str(exc))
        return ""


async def extract_text_from_docx(file_path: str) -> str:
    """
    Extract plain text from a DOCX file using python-docx.

    Extracts paragraph text and table cell text, preserving paragraph breaks.
    """
    try:
        import docx  # python-docx

        doc = docx.Document(file_path)
        parts: list[str] = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())

        full_text = "\n\n".join(parts)
        log.info(
            "extract_text_from_docx.done",
            file=file_path,
            paragraphs=len(doc.paragraphs),
            chars=len(full_text),
        )
        return full_text
    except Exception as exc:
        log.error("extract_text_from_docx.failed", file=file_path, error=str(exc))
        return ""


async def extract_text_from_image(file_path: str) -> str:
    """
    Extract text from an image file using Azure Document Intelligence.

    Falls back to empty string if the service is unavailable.
    """
    try:
        from azure.ai.documentintelligence import DocumentIntelligenceClient
        from azure.core.credentials import AzureKeyCredential

        from app.config import settings

        client = DocumentIntelligenceClient(
            endpoint=settings.DOCINTEL_ENDPOINT,
            credential=AzureKeyCredential(settings.DOCINTEL_KEY),
        )

        with open(file_path, "rb") as f:
            file_bytes = f.read()

        poller = client.begin_analyze_document(
            "prebuilt-read",
            analyze_request=file_bytes,
            content_type="application/octet-stream",
        )
        result = poller.result()
        text = result.content or ""
        log.info("extract_text_from_image.done", file=file_path, chars=len(text))
        return text
    except ImportError:
        log.warning("extract_text_from_image.no_sdk")
        return ""
    except Exception as exc:
        log.error("extract_text_from_image.failed", file=file_path, error=str(exc))
        return ""


# ---------------------------------------------------------------------------
# AI document analysis
# ---------------------------------------------------------------------------


async def analyze_document_with_ai(
    text: str,
    doc_type: str,
    case_context: str | None = None,
) -> dict[str, Any]:
    """
    Perform AI analysis on extracted document text.

    Returns a structured dict:
        {summary, key_dates, key_persons, legal_implications,
         action_required, urgency}
    """
    log.info(
        "analyze_document_with_ai.start",
        doc_type=doc_type,
        text_len=len(text),
    )

    context_block = (
        f"\n**Fallkontext:**\n{case_context}\n" if case_context else ""
    )

    system = (
        "Du bist ein erfahrener Rechtsassistent in Deutschland. "
        "Analysiere Rechtsdokumente präzise und strukturiert. "
        "Antworte ausschließlich als valides JSON-Objekt ohne Markdown-Codeblöcke."
    )

    user_prompt = f"""Analysiere das folgende {doc_type}-Dokument.{context_block}

Gib ein JSON-Objekt mit diesen Feldern zurück:
{{
  "summary": "Kurze Zusammenfassung des Dokuments (2-4 Sätze)",
  "key_dates": [
    {{"date": "TT.MM.JJJJ oder Beschreibung", "significance": "Bedeutung dieses Datums"}}
  ],
  "key_persons": [
    {{"name": "Name/Institution", "role": "Rolle im Dokument"}}
  ],
  "legal_implications": "Rechtliche Bedeutung, Konsequenzen und relevante Paragraphen",
  "action_required": "Konkrete Maßnahmen die der Nutzer ergreifen sollte",
  "urgency": "low|medium|high|critical",
  "document_category": "police_report|court_letter|contract|id_document|medical|financial|correspondence|other"
}}

**Dokumenttext:**
\"\"\"{text[:6000]}\"\"\"
"""

    messages = [
        {"role": "system", "content": system},
        {"role": "user", "content": user_prompt},
    ]

    raw = await chat_completion(messages, temperature=0.1, max_tokens=2000)
    assert isinstance(raw, str)

    try:
        cleaned = _strip_fences(raw)
        result = json.loads(cleaned)
        log.info("analyze_document_with_ai.done", urgency=result.get("urgency"))
        return result
    except (json.JSONDecodeError, Exception) as exc:
        log.warning(
            "analyze_document_with_ai.json_failed",
            error=str(exc),
            raw=raw[:300],
        )
        return {
            "summary": raw[:500],
            "key_dates": [],
            "key_persons": [],
            "legal_implications": "Analyse konnte nicht strukturiert werden.",
            "action_required": "Dokument manuell prüfen und Anwalt konsultieren.",
            "urgency": "medium",
            "document_category": "other",
        }


# ---------------------------------------------------------------------------
# Full document processing pipeline
# ---------------------------------------------------------------------------


async def process_uploaded_document(
    document: Document,
    db: AsyncSession,
) -> Document:
    """
    Full processing pipeline for an uploaded document:

    1. Detect file type from mime_type
    2. Extract text
    3. Run AI analysis
    4. Update the document record in the database

    Sets processing_status to "completed" on success, "failed" on error.
    """
    log.info(
        "process_uploaded_document.start",
        doc_id=str(document.id),
        mime=document.mime_type,
    )

    document.processing_status = "processing"
    await db.flush()

    try:
        # 1. Extract text
        extracted_text = ""
        mime = document.mime_type.lower()

        if "pdf" in mime:
            extracted_text = await extract_text_from_pdf(document.file_path)
        elif "docx" in mime or "openxmlformats" in mime or "word" in mime:
            extracted_text = await extract_text_from_docx(document.file_path)
        elif any(img in mime for img in ("jpeg", "jpg", "png", "tiff", "bmp", "gif")):
            extracted_text = await extract_text_from_image(document.file_path)
        elif "text" in mime or "plain" in mime:
            async with aiofiles.open(document.file_path, mode="r", encoding="utf-8", errors="replace") as f:
                extracted_text = await f.read()
        else:
            log.warning(
                "process_uploaded_document.unsupported_mime",
                mime=document.mime_type,
            )

        document.extracted_text = extracted_text

        # 2. AI analysis (only if we have text)
        if extracted_text.strip():
            analysis_result = await analyze_document_with_ai(
                text=extracted_text,
                doc_type=document.document_category,
            )

            # Store as JSON string in the ai_analysis column
            document.ai_analysis = json.dumps(analysis_result, ensure_ascii=False)

            # Update category if AI detected a better one
            detected_category = analysis_result.get("document_category")
            if detected_category and detected_category != "other":
                document.document_category = detected_category
        else:
            document.ai_analysis = json.dumps(
                {"summary": "Kein Text extrahiert.", "urgency": "low"},
                ensure_ascii=False,
            )

        document.processing_status = "completed"
        log.info("process_uploaded_document.done", doc_id=str(document.id))

    except Exception as exc:
        log.error(
            "process_uploaded_document.failed",
            doc_id=str(document.id),
            error=str(exc),
        )
        document.processing_status = "failed"
        document.ai_analysis = json.dumps(
            {"error": str(exc), "summary": "Verarbeitung fehlgeschlagen."},
            ensure_ascii=False,
        )

    await db.flush()
    await db.refresh(document)
    return document


# ---------------------------------------------------------------------------
# Filename utilities
# ---------------------------------------------------------------------------


def get_safe_filename(filename: str) -> str:
    """
    Sanitise a filename to be safe for storage on the filesystem.

    - Normalise unicode (NFKD → ASCII)
    - Replace spaces with underscores
    - Remove all characters that are not alphanumeric, dashes, dots, or underscores
    - Prepend a short UUID to guarantee uniqueness
    """
    # Normalise unicode
    normalised = unicodedata.normalize("NFKD", filename)
    ascii_name = normalised.encode("ascii", "ignore").decode("ascii")

    # Keep extension
    path = Path(ascii_name)
    stem = re.sub(r"[^\w\-.]", "_", path.stem).strip("._")
    suffix = path.suffix.lower()

    # Guard against empty stems
    if not stem:
        stem = "document"

    unique_prefix = str(uuid.uuid4())[:8]
    return f"{unique_prefix}_{stem}{suffix}"


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _strip_fences(text: str) -> str:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        parts = cleaned.split("```")
        cleaned = parts[1] if len(parts) > 1 else cleaned
        if cleaned.startswith("json"):
            cleaned = cleaned[4:].strip()
    return cleaned
